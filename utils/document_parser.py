"""
Rezyume va hujjatlardan matn ajratib olish (PDF, DOCX, TXT)
"""

import os
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str | Path) -> str:
    """PDF fayldan barcha sahifalar matnini ajratib oladi."""
    text_chunks = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        for idx, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_chunks.append(page_text.strip())
        
        extracted = "\n\n".join(text_chunks)
        return extracted.strip()
    except Exception as e:
        logger.error(f"PDF matnini o'qishda xatolik: {e}")
        return ""

def extract_text_from_docx(file_path: str | Path) -> str:
    """DOCX fayldan matn va jadvallarni ajratib oladi."""
    text_chunks = []
    try:
        import docx
        doc = docx.Document(str(file_path))
        
        # Paragraflarni o'qish
        for para in doc.paragraphs:
            if para.text.strip():
                text_chunks.append(para.text.strip())
        
        # Jadvallardagi matnlarni o'qish
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_chunks.append(" | ".join(row_text))
                    
        extracted = "\n".join(text_chunks)
        return extracted.strip()
    except Exception as e:
        logger.error(f"DOCX matnini o'qishda xatolik: {e}")
        return ""

def extract_text_from_file(file_path: str | Path) -> str:
    """Fayl kengaytmasiga qarab tegishli parserni ishga tushiradi."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    elif suffix in [".docx", ".doc"]:
        return extract_text_from_docx(path)
    elif suffix in [".txt", ".md"]:
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception as e:
            logger.error(f"Matnli faylni o'qishda xatolik: {e}")
            return ""
    else:
        return ""
